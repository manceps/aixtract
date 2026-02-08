"""Word document converter."""
from __future__ import annotations

import io
from pathlib import Path
from typing import BinaryIO, ClassVar

from aixtract.converters.base import BaseConverter
from aixtract.core.registry import ConverterRegistry
from aixtract.models.result import DocumentMetadata, ExtractionResult


@ConverterRegistry.register
class DOCXConverter(BaseConverter):
    """Extract text and metadata from Word documents."""

    name: ClassVar[str] = "docx"
    supported_extensions: ClassVar[tuple[str, ...]] = (".docx", ".doc")
    supported_mimetypes: ClassVar[tuple[str, ...]] = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    )
    requires: ClassVar[tuple[str, ...]] = ("docx",)

    def extract(
        self,
        source: Path | BinaryIO | bytes,
        filename: str | None = None,
    ) -> ExtractionResult:
        """Extract content from Word document."""
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError

        content_bytes, file_path = self._read_source(source)

        try:
            doc = Document(io.BytesIO(content_bytes))
        except PackageNotFoundError:
            return ExtractionResult(
                success=False,
                error="Invalid or corrupted Word document",
                metadata=DocumentMetadata(
                    filename=filename or "document.docx",
                    file_path=file_path,
                ),
            )

        # Extract metadata
        props = doc.core_properties
        metadata = DocumentMetadata(
            filename=filename or (file_path.name if file_path else "document.docx"),
            file_path=file_path,
            format_detected="docx",
            title=props.title,
            author=props.author,
            subject=props.subject,
            keywords=props.keywords.split(",") if props.keywords else [],
            created_at=props.created,
            modified_at=props.modified,
            converter_used=self.name,
        )

        # Extract content
        text_parts = []
        markdown_parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            text_parts.append(text)

            # Apply markdown formatting based on style
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                level = int(style_name[-1]) if style_name[-1].isdigit() else 1
                markdown_parts.append("#" * level + " " + text)
            else:
                markdown_parts.append(text)

        # Extract tables
        for table in doc.tables:
            markdown_parts.append(self._table_to_markdown(table))

        content = "\n\n".join(text_parts)
        content_markdown = "\n\n".join(markdown_parts)

        metadata.word_count = len(content.split())
        metadata.char_count = len(content)

        return ExtractionResult(
            success=True,
            content=content,
            content_markdown=content_markdown,
            metadata=metadata,
        )

    def _table_to_markdown(self, table) -> str:
        """Convert DOCX table to markdown."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)

        if not rows:
            return ""

        lines = []
        # Header
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")

        # Data rows
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")

        return "\n".join(lines)
